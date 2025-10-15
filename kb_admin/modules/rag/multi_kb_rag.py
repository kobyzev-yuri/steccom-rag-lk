"""
Multi Knowledge Base RAG System
–°–∏—Å—Ç–µ–º–∞ RAG –¥–ª—è —Ä–∞–±–æ—Ç—ã —Å –Ω–µ—Å–∫–æ–ª—å–∫–∏–º–∏ –±–∞–∑–∞–º–∏ –∑–Ω–∞–Ω–∏–π
"""

import streamlit as st
import json
import sqlite3
from typing import List, Dict, Optional, Tuple
from pathlib import Path
import numpy as np
from langchain_community.embeddings import OllamaEmbeddings
try:
    from langchain_huggingface import HuggingFaceEmbeddings
    HUGGINGFACE_AVAILABLE = True
except ImportError:
    HUGGINGFACE_AVAILABLE = False
    print("‚ö†Ô∏è HuggingFaceEmbeddings –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω: langchain_huggingface –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.chat_models import ChatOllama
from langchain_openai import ChatOpenAI
import os
from datetime import datetime
from langchain.prompts import ChatPromptTemplate
from langchain.schema.runnable import RunnablePassthrough
from langchain.schema import StrOutputParser
from langchain.schema import Document

class MultiKBRAG:
    def __init__(self, db_path: str = None,
                 chat_provider: Optional[str] = None,
                 chat_model: Optional[str] = None,
                 proxy_base_url: Optional[str] = None,
                 proxy_api_key: Optional[str] = None,
                 temperature: float = 0.2):
        if db_path is None:
            # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∞–±—Å–æ–ª—é—Ç–Ω—ã–π –ø—É—Ç—å –∫ –ë–î –±–∞–∑ –∑–Ω–∞–Ω–∏–π
            current_dir = Path(__file__).parent
            # –ò–¥–µ–º –Ω–∞ 3 —É—Ä–æ–≤–Ω—è –≤–≤–µ—Ä—Ö: modules/rag -> modules -> kb_admin -> steccom
            self.db_path = current_dir.parent.parent.parent / "kb_admin" / "kbs.db"
            # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –≤ –∞–±—Å–æ–ª—é—Ç–Ω—ã–π –ø—É—Ç—å
            self.db_path = self.db_path.resolve()
            
            # –ï—Å–ª–∏ –ë–î –Ω–µ –Ω–∞–π–¥–µ–Ω–∞, –ø—Ä–æ–±—É–µ–º –ø—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ –∫–æ—Ä–Ω—è –ø—Ä–æ–µ–∫—Ç–∞
            if not self.db_path.exists():
                root_dir = current_dir.parent.parent.parent.parent  # <repo>
                self.db_path = root_dir / "kb_admin" / "kbs.db"
                self.db_path = self.db_path.resolve()
        else:
            self.db_path = db_path
        # Embeddings backend: default to multilingual HF model if configured
        embedding_provider = os.getenv("EMBEDDING_PROVIDER", "ollama").lower()
        if embedding_provider in ("hf", "huggingface") and HUGGINGFACE_AVAILABLE:
            embedding_model = os.getenv("EMBEDDING_MODEL", "intfloat/multilingual-e5-base")
            self.embeddings = HuggingFaceEmbeddings(model_name=embedding_model, 
                                                    model_kwargs={"device": "cpu"},
                                                    encode_kwargs={"normalize_embeddings": True})
            self._embedding_backend = {"provider": "huggingface", "model": embedding_model}
        else:
            # –ò—Å–ø–æ–ª—å–∑—É–µ–º HuggingFace –¥–ª—è —ç–º–±–µ–¥–¥–∏–Ω–≥–æ–≤ (–±—ã—Å—Ç—Ä–µ–µ –∏ —ç—Ñ—Ñ–µ–∫—Ç–∏–≤–Ω–µ–µ)
            embedding_model = os.getenv("EMBEDDING_MODEL", "intfloat/multilingual-e5-base")
            self.embeddings = HuggingFaceEmbeddings(model_name=embedding_model, 
                                                    model_kwargs={"device": "cpu"},
                                                    encode_kwargs={"normalize_embeddings": True})
            self._embedding_backend = {"provider": "huggingface", "model": embedding_model}
        # Resolve chat backend configuration (constructor overrides env)
        env_use_proxy = os.getenv("USE_PROXYAPI", "false").lower() == "true"
        resolved_provider = (chat_provider or ("proxyapi" if env_use_proxy else "ollama")).lower()
        resolved_temperature = float(os.getenv("PROXYAPI_TEMPERATURE", str(temperature)))
        if resolved_provider == "proxyapi":
            resolved_base_url = proxy_base_url or os.getenv("PROXYAPI_BASE_URL", "https://api.proxyapi.ru/openai/v1")
            # Support OPEN_AI_API_KEY as proxyapi key alias
            resolved_api_key = proxy_api_key or os.getenv("PROXYAPI_API_KEY") or os.getenv("OPEN_AI_API_KEY", "")
            resolved_model = chat_model or os.getenv("PROXYAPI_CHAT_MODEL", "gpt-4o")
            self.chat_model = ChatOpenAI(
                model=resolved_model,
                openai_api_key=resolved_api_key,
                base_url=resolved_base_url,
                temperature=resolved_temperature
            )
            self._chat_backend = {
                'provider': 'proxyapi', 'model': resolved_model,
                'base_url': resolved_base_url, 'temperature': resolved_temperature
            }
        elif resolved_provider == "openai":
            # Direct OpenAI usage (no custom base_url)
            resolved_model = chat_model or os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")
            openai_key = os.getenv("OPENAI_API_KEY", "")
            self.chat_model = ChatOpenAI(
                model=resolved_model,
                openai_api_key=openai_key,
                temperature=resolved_temperature
            )
            self._chat_backend = {'provider': 'openai', 'model': resolved_model, 'temperature': resolved_temperature}
        else:
            # –¢–æ–ª—å–∫–æ ProxyAPI –ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ—Ç—Å—è
            raise Exception("ProxyAPI –Ω–µ –Ω–∞—Å—Ç—Ä–æ–µ–Ω. –£—Å—Ç–∞–Ω–æ–≤–∏—Ç–µ PROXYAPI_KEY –≤ config.env")
        self.vectorstores = {}  # kb_id -> vectorstore
        self.kb_metadata = {}   # kb_id -> metadata
        self.kb_chunks = {}     # kb_id -> List[Document]
        
        # –ó–∞–≥—Ä—É–∂–∞–µ–º –≤—Å–µ –∞–∫—Ç–∏–≤–Ω—ã–µ –±–∞–∑—ã –∑–Ω–∞–Ω–∏–π –ø—Ä–∏ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏
        try:
            self.load_all_active_kbs()
        except Exception as e:
            print(f"‚ö†Ô∏è –ù–µ —É–¥–∞–ª–æ—Å—å –∑–∞–≥—Ä—É–∑–∏—Ç—å –±–∞–∑—ã –∑–Ω–∞–Ω–∏–π: {e}")

    def get_chat_backend_info(self) -> Dict:
        """–í–æ–∑–≤—Ä–∞—â–∞–µ—Ç —Ç–µ–∫—É—â—É—é –∫–æ–Ω—Ñ–∏–≥—É—Ä–∞—Ü–∏—é LLM –ø—Ä–æ–≤–∞–π–¥–µ—Ä–∞ –¥–ª—è RAG."""
        try:
            return dict(self._chat_backend)
        except Exception:
            return {"provider": "unknown"}

    def set_chat_backend(self, provider: str, model: str,
                         base_url: Optional[str] = None,
                         api_key: Optional[str] = None,
                         temperature: float = 0.2) -> None:
        """Reconfigure chat backend without touching vector stores."""
        provider = (provider or "ollama").lower()
        if provider == "proxyapi":
            self.chat_model = ChatOpenAI(
                model=model,
                openai_api_key=api_key or os.getenv("PROXYAPI_API_KEY") or os.getenv("OPEN_AI_API_KEY", ""),
                base_url=base_url or os.getenv("PROXYAPI_BASE_URL", "https://api.proxyapi.ru/openai/v1"),
                temperature=temperature
            )
            self._chat_backend = {
                'provider': 'proxyapi', 'model': model,
                'base_url': base_url, 'temperature': temperature
            }
        elif provider == "openai":
            self.chat_model = ChatOpenAI(
                model=model or os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini"),
                openai_api_key=api_key or os.getenv("OPENAI_API_KEY", ""),
                temperature=temperature
            )
            self._chat_backend = {'provider': 'openai', 'model': model, 'temperature': temperature}
        else:
            self.chat_model = ChatOllama(model=model or os.getenv("OLLAMA_CHAT_MODEL", "qwen3:8b"))
            self._chat_backend = {'provider': 'ollama', 'model': model}

    def _ensure_usage_table(self) -> None:
        """Create llm_usage table if not exists."""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(
            """
            CREATE TABLE IF NOT EXISTS llm_usage (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT NOT NULL,
                provider TEXT,
                model TEXT,
                prompt_tokens INTEGER,
                completion_tokens INTEGER,
                total_tokens INTEGER,
                question TEXT,
                response_length INTEGER
            )
            """
        )
        conn.commit()
        conn.close()

    def _log_llm_usage(self, provider: Optional[str], model: Optional[str],
                        prompt_tokens: Optional[int], completion_tokens: Optional[int],
                        total_tokens: Optional[int], question: str,
                        response_length: int) -> None:
        """Insert usage row into llm_usage table."""
        self._ensure_usage_table()
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        c.execute(
            """
            INSERT INTO llm_usage (
                timestamp, provider, model, prompt_tokens, completion_tokens, total_tokens, question, response_length
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                datetime.utcnow().isoformat(),
                provider, model,
                prompt_tokens if isinstance(prompt_tokens, int) else None,
                completion_tokens if isinstance(completion_tokens, int) else None,
                total_tokens if isinstance(total_tokens, int) else None,
                question, int(response_length)
            )
        )
        conn.commit()
        conn.close()
        self.vectorstores = {}  # kb_id -> vectorstore
        self.kb_metadata = {}   # kb_id -> metadata
        
    def load_knowledge_base(self, kb_id: int) -> bool:
        """Load specific knowledge base into memory"""
        try:
            # Try to load existing vectorstore first
            # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∞–±—Å–æ–ª—é—Ç–Ω—ã–π –ø—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ –ø–∞–ø–∫–∏ kb_admin
            current_dir = Path(__file__).parent.parent.parent  # <repo>/kb_admin
            # –ï–¥–∏–Ω—ã–π –ø—É—Ç—å —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏—è/–∑–∞–≥—Ä—É–∑–∫–∏ –∏–Ω–¥–µ–∫—Å–∞
            vectorstore_path = current_dir / "data" / "knowledge_bases" / f"vectorstore_{kb_id}"
            
            # –ï—Å–ª–∏ –ø—É—Ç—å –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç, –ø—Ä–æ–±—É–µ–º –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ –∫–æ—Ä–Ω—è –ø—Ä–æ–µ–∫—Ç–∞
            if not Path(vectorstore_path).exists():
                # –ü—Ä–æ–±—É–µ–º –ø—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ –∫–æ—Ä–Ω—è –ø—Ä–æ–µ–∫—Ç–∞ (–∫–æ–≥–¥–∞ –∑–∞–ø—É—Å–∫–∞–µ–º –∏–∑ –∫–æ—Ä–Ω—è)
                root_dir = Path(__file__).parent.parent.parent.parent  # <repo>
                vectorstore_path = root_dir / "kb_admin" / "data" / "knowledge_bases" / f"vectorstore_{kb_id}"
                print(f"üîÑ –ü—Ä–æ–±—É–µ–º –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã–π –ø—É—Ç—å: {vectorstore_path}")
            if Path(vectorstore_path).exists():
                try:
                    vectorstore = FAISS.load_local(str(vectorstore_path), self.embeddings, allow_dangerous_deserialization=True)
                    # Get KB info from database
                    conn = sqlite3.connect(self.db_path)
                    c = conn.cursor()
                    c.execute("SELECT * FROM knowledge_bases WHERE id = ? AND is_active = 1", (kb_id,))
                    kb_info = c.fetchone()
                    
                    if kb_info:
                        self.vectorstores[kb_id] = vectorstore
                        
                        # –ü–æ–¥—Å—á–∏—Ç—ã–≤–∞–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –∏ —á–∞–Ω–∫–æ–≤
                        c.execute("SELECT COUNT(*) FROM knowledge_documents WHERE kb_id = ? AND processed = 1", (kb_id,))
                        doc_count = c.fetchone()[0]
                        
                        c.execute("SELECT COUNT(*) FROM document_chunks dc JOIN knowledge_documents kd ON dc.doc_id = kd.id WHERE kd.kb_id = ? AND kd.processed = 1", (kb_id,))
                        chunk_count = c.fetchone()[0]
                        
                        conn.close()
                        
                        self.kb_metadata[kb_id] = {
                            'name': kb_info[1],
                            'description': kb_info[2],
                            'category': kb_info[3],
                            'doc_count': doc_count,
                            'chunk_count': chunk_count
                        }
                        print(f"‚úÖ –ó–∞–≥—Ä—É–∂–µ–Ω —Å–æ—Ö—Ä–∞–Ω–µ–Ω–Ω—ã–π –≤–µ–∫—Ç–æ—Ä–Ω—ã–π –∏–Ω–¥–µ–∫—Å –¥–ª—è KB {kb_id} (–¥–æ–∫—É–º–µ–Ω—Ç–æ–≤: {doc_count}, —á–∞–Ω–∫–æ–≤: {chunk_count})")
                        return True
                except Exception as e:
                    print(f"‚ö†Ô∏è –û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –≤–µ–∫—Ç–æ—Ä–Ω–æ–≥–æ –∏–Ω–¥–µ–∫—Å–∞ –¥–ª—è KB {kb_id}: {e}")
            
            # Get KB info from database
            conn = sqlite3.connect(self.db_path)
            c = conn.cursor()
            
            c.execute("SELECT * FROM knowledge_bases WHERE id = ? AND is_active = 1", (kb_id,))
            kb_info = c.fetchone()
            
            if not kb_info:
                conn.close()
                return False
            
            # Get documents for this KB
            c.execute('''
                SELECT * FROM knowledge_documents 
                WHERE kb_id = ? AND processed = 1
                ORDER BY upload_date DESC
            ''', (kb_id,))
            
            documents = c.fetchall()
            
            # Get images for this KB (only if table exists and has data)
            try:
                c.execute('''
                    SELECT * FROM knowledge_images 
                    WHERE kb_id = ? AND processed = 1
                    ORDER BY upload_date DESC
                ''', (kb_id,))
                images = c.fetchall()
            except sqlite3.OperationalError as e:
                if "no such table" in str(e).lower():
                    print(f"‚ö†Ô∏è –¢–∞–±–ª–∏—Ü–∞ knowledge_images –Ω–µ –Ω–∞–π–¥–µ–Ω–∞ –¥–ª—è KB {kb_id}, –ø—Ä–æ–ø—É—Å–∫–∞–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è")
                    images = []
                else:
                    raise
            conn.close()
            
            if not documents and not images:
                return False
            
            # Process documents into vectorstore
            documents_list = []
            for doc in documents:
                # Load document content
                doc_path = doc[3]  # file_path column
                doc_meta_raw = doc[9]  # metadata column (JSON)
                meta_abstract = None
                try:
                    if isinstance(doc_meta_raw, str) and doc_meta_raw.strip():
                        import json as _json
                        _m = _json.loads(doc_meta_raw)
                        meta_abstract = _m.get('abstract') or _m.get('summary')
                except Exception:
                    meta_abstract = None
                if doc_path and Path(doc_path).exists():
                    content = ""
                    extraction_error = None
                    suffix = Path(doc_path).suffix.lower()
                    if suffix == ".json":
                        # JSON KB: —Å–æ–±—Ä–∞—Ç—å —Ç–µ–∫—Å—Ç –∏–∑ –ø–æ–ª–µ–π title/content
                        try:
                            import json as _json
                            with open(doc_path, 'r', encoding='utf-8') as f:
                                data = _json.load(f)
                            parts = []
                            # –ø–æ–¥–¥–µ—Ä–∂–∫–∞ —Ñ–æ—Ä–º–∞—Ç–∞: —Å–ø–∏—Å–æ–∫ –æ–±—ä–µ–∫—Ç–æ–≤ —Å –ø–æ–ª—è–º–∏ title, content(list of {title,text})
                            if isinstance(data, list):
                                for item in data:
                                    try:
                                        it_title = str(item.get('title', ''))
                                        parts.append(it_title)
                                        it_content = item.get('content', [])
                                        if isinstance(it_content, list):
                                            for c in it_content:
                                                t = c.get('title')
                                                x = c.get('text')
                                                if t:
                                                    parts.append(str(t))
                                                if x:
                                                    parts.append(str(x))
                                    except Exception:
                                        continue
                            else:
                                # –µ—Å–ª–∏ –æ–±—ä–µ–∫—Ç, –ø—Ä–æ—Å—Ç–æ —Å–µ—Ä–∏–∞–ª–∏–∑—É–µ–º
                                parts.append(_json.dumps(data, ensure_ascii=False))
                            content = "\n".join([p for p in parts if p]).strip()
                        except Exception as e:
                            extraction_error = f"json-extract: {e}"
                    else:
                        # Handle common text-bearing formats
                        try:
                            if suffix in (".txt", ".md"):
                                with open(doc_path, 'r', encoding='utf-8', errors='ignore') as f:
                                    content = f.read()
                            elif suffix in (".docx", ".doc"):
                                try:
                                    from docx import Document as _DocxDocument
                                    docx_doc = _DocxDocument(doc_path)
                                    parts = []
                                    for paragraph in docx_doc.paragraphs:
                                        if paragraph.text and paragraph.text.strip():
                                            parts.append(paragraph.text.strip())
                                    for table in docx_doc.tables:
                                        for row in table.rows:
                                            row_text = []
                                            for cell in row.cells:
                                                if cell.text and cell.text.strip():
                                                    row_text.append(cell.text.strip())
                                            if row_text:
                                                parts.append(" | ".join(row_text))
                                    content = "\n".join(parts)
                                except Exception as _docx_err:
                                    extraction_error = f"docx-extract: {_docx_err}"
                            else:
                                # Assume PDF or other binary with text
                                import PyPDF2
                                with open(doc_path, 'rb') as file:
                                    pdf_reader = PyPDF2.PdfReader(file)
                                    for page in pdf_reader.pages:
                                        extracted = page.extract_text() or ""
                                        content += extracted + "\n"
                        except Exception as e:
                            extraction_error = str(e)

                        # Fallbacks for PDFs when text is too small
                        if suffix == ".pdf" and len(content.strip()) < 50:
                            try:
                                import fitz  # PyMuPDF
                                with fitz.open(doc_path) as pdf_doc:
                                    content = "\n".join([page.get_text() or "" for page in pdf_doc])
                            except Exception as e:
                                extraction_error = extraction_error or str(e)

                    # If still empty, create minimal placeholder
                    if not content.strip():
                        content = f"–î–æ–∫—É–º–µ–Ω—Ç: {doc[2]} (—Ñ–∞–π–ª: {Path(doc_path).name}) ‚Äî —Ç–µ–∫—Å—Ç –Ω–µ –∏–∑–≤–ª–µ—á–µ–Ω."

                    # Prepend abstract from metadata if available to boost relevance
                    if isinstance(meta_abstract, str) and meta_abstract.strip():
                        content = f"ABSTRACT (–∏–∑ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö):\n{meta_abstract.strip()}\n\n{content}"

                    # Append document regardless so KB can load
                    doc_metadata = {
                        'kb_id': kb_id,
                        'doc_id': doc[0],
                        'title': doc[2],
                        'source': doc_path,
                        'category': kb_info[3]
                    }
                    if extraction_error:
                        doc_metadata['extraction_error'] = extraction_error

                    documents_list.append(Document(
                        page_content=content,
                        metadata=doc_metadata
                    ))
            
            # Process images into documents
            for image in images:
                image_id = image[0]
                doc_id = image[2]  # doc_id column
                image_path = image[3]  # image_path column
                image_name = image[4]  # image_name column
                image_description = image[5]  # image_description column
                llava_analysis = image[7]  # llava_analysis column
                
                # –°–æ–∑–¥–∞–µ–º –∫–æ–Ω—Ç–µ–Ω—Ç –¥–ª—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                image_content = f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ: {image_name}\n"
                if image_description:
                    image_content += f"–û–ø–∏—Å–∞–Ω–∏–µ: {image_description}\n"
                if llava_analysis:
                    image_content += f"–ê–Ω–∞–ª–∏–∑ LLaVA: {llava_analysis}\n"
                
                # –°–æ–∑–¥–∞–µ–º –¥–æ–∫—É–º–µ–Ω—Ç –¥–ª—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
                image_doc = Document(
                    page_content=image_content,
                    metadata={
                        'kb_id': kb_id,
                        'image_id': image_id,
                        'doc_id': doc_id,
                        'title': f"–ò–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ: {image_name}",
                        'source': image_path,
                        'category': kb_info[3],
                        'content_type': 'image',
                        'image_name': image_name,
                        'image_description': image_description,
                        'llava_analysis': llava_analysis
                    }
                )
                documents_list.append(image_doc)
            
            if not documents_list:
                # Should not happen, but guard: create one placeholder to allow KB to load
                documents_list.append(Document(
                    page_content=f"–ë–∞–∑–∞ –∑–Ω–∞–Ω–∏–π {kb_info[1]} –Ω–µ —Å–æ–¥–µ—Ä–∂–∏—Ç –¥–æ—Å—Ç—É–ø–Ω—ã—Ö –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ –∏–ª–∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π.",
                    metadata={
                        'kb_id': kb_id,
                        'doc_id': -1,
                        'title': '–ü—É—Å—Ç–∞—è –±–∞–∑–∞ –∑–Ω–∞–Ω–∏–π',
                        'source': 'db',
                        'category': kb_info[3]
                    }
                ))
            
            # Adaptive chunking per document
            def _pick_splitter_for(text: str) -> RecursiveCharacterTextSplitter:
                t = (text or "").lower()
                # Heuristic: Q/A style if many question marks and bullets
                qa_like = (t.count('?') >= 3 and ('‚Ä¢' in text or '- ' in text)) or ('–≤–æ–ø—Ä–æ—Å' in t and '–æ—Ç–≤–µ—Ç' in t) or ('question' in t and 'answer' in t)
                # Heuristic: code-like if fenced blocks or many braces/semicolons/indentation
                code_like = ('```' in text) or (t.count('{') + t.count('}') > 10) or (t.count(';') > 20)
                if qa_like:
                    return RecursiveCharacterTextSplitter(
                        chunk_size=800,
                        chunk_overlap=120,
                        length_function=len,
                        separators=["\n\n", "\n", "? ", "?\n", "‚Ä¢ ", "- ", ". ", " ", ""]
                    )
                if code_like:
                    return RecursiveCharacterTextSplitter(
                        chunk_size=600,
                        chunk_overlap=80,
                        length_function=len,
                        separators=["\n\n", "\n", "; ", ";\n", ") ", ": ", ". ", " ", ""]
                    )
                return RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len,
                    separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
                )

            chunks = []
            for d in documents_list:
                splitter = _pick_splitter_for(d.page_content)
                chunks.extend(splitter.split_documents([d]))
            
            # Create vectorstore
            vectorstore = FAISS.from_documents(chunks, self.embeddings)
            
            # Try to save vectorstore to disk for faster loading
            try:
                # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∞–±—Å–æ–ª—é—Ç–Ω—ã–π –ø—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ –ø–∞–ø–∫–∏ kb_admin
                current_dir = Path(__file__).parent.parent.parent
                vectorstore_path = current_dir / "data" / "knowledge_bases" / f"vectorstore_{kb_id}"
                
                # –ï—Å–ª–∏ –ø—É—Ç—å –Ω–µ —Å—É—â–µ—Å—Ç–≤—É–µ—Ç, –ø—Ä–æ–±—É–µ–º –ø—É—Ç—å –æ—Ç–Ω–æ—Å–∏—Ç–µ–ª—å–Ω–æ –∫–æ—Ä–Ω—è –ø—Ä–æ–µ–∫—Ç–∞
                if not vectorstore_path.parent.exists():
                    root_dir = Path(__file__).parent.parent.parent.parent  # <repo>
                    vectorstore_path = root_dir / "kb_admin" / "data" / "knowledge_bases" / f"vectorstore_{kb_id}"
                
                vectorstore.save_local(str(vectorstore_path))
            except Exception as e:
                print(f"‚ö†Ô∏è –ù–µ —É–¥–∞–ª–æ—Å—å —Å–æ—Ö—Ä–∞–Ω–∏—Ç—å –≤–µ–∫—Ç–æ—Ä–Ω—ã–π –∏–Ω–¥–µ–∫—Å –¥–ª—è KB {kb_id}: {e}")
            
            # Store in memory
            self.vectorstores[kb_id] = vectorstore
            self.kb_metadata[kb_id] = {
                'name': kb_info[1],
                'description': kb_info[2],
                'category': kb_info[3],
                'doc_count': len(documents_list),
                'chunk_count': len(chunks)
            }
            self.kb_chunks[kb_id] = chunks
            
            return True
            
        except Exception as e:
            st.error(f"–û—à–∏–±–∫–∞ –∑–∞–≥—Ä—É–∑–∫–∏ –±–∞–∑—ã –∑–Ω–∞–Ω–∏–π {kb_id}: {e}")
            return False
    
    def load_all_active_kbs(self) -> int:
        """Load all active knowledge bases"""
        conn = sqlite3.connect(self.db_path)
        c = conn.cursor()
        
        c.execute("SELECT id FROM knowledge_bases WHERE is_active = 1")
        kb_ids = [row[0] for row in c.fetchall()]
        conn.close()
        
        loaded_count = 0
        for kb_id in kb_ids:
            if self.load_knowledge_base(kb_id):
                loaded_count += 1
        
        return loaded_count
    
    def get_available_kbs(self) -> List[Dict]:
        """Get list of available knowledge bases"""
        return [
            {
                'kb_id': kb_id,
                **metadata
            }
            for kb_id, metadata in self.kb_metadata.items()
        ]
    
    def search_across_kbs(self, query: str, kb_ids: List[int] = None, k: int = 5) -> List[Document]:
        """Search across multiple knowledge bases with hybrid search"""
        if kb_ids is None:
            kb_ids = list(self.vectorstores.keys())
        
        all_results = []
        
        for kb_id in kb_ids:
            if kb_id in self.vectorstores:
                # Vector search
                vectorstore = self.vectorstores[kb_id]
                vector_results = vectorstore.similarity_search(query, k=k)
                
                # Add KB metadata to results
                for doc in vector_results:
                    doc.metadata['kb_name'] = self.kb_metadata[kb_id]['name']
                    doc.metadata['kb_category'] = self.kb_metadata[kb_id]['category']
                
                all_results.extend(vector_results)
                
                # Text search for exact matches
                text_results = self._text_search_in_kb(query, kb_id, k)
                all_results.extend(text_results)
        
        # Remove duplicates and sort by relevance
        unique_results = self._deduplicate_documents(all_results)
        return unique_results[:k*len(kb_ids)]
    
    def search_in_kb(self, query: str, kb_id: int, k: int = 5) -> List[Document]:
        """Search within specific knowledge base"""
        if kb_id not in self.vectorstores:
            return []
        
        vectorstore = self.vectorstores[kb_id]
        results = vectorstore.similarity_search(query, k=k)
        
        # Add KB metadata
        for doc in results:
            doc.metadata['kb_name'] = self.kb_metadata[kb_id]['name']
            doc.metadata['kb_category'] = self.kb_metadata[kb_id]['category']
        
        return results
    
    def _text_search_in_kb(self, query: str, kb_id: int, k: int = 5) -> List[Document]:
        """Strong keyword search over stored chunks within specific knowledge base"""
        if kb_id not in self.kb_chunks:
            return []
        
        query_lower = query.lower()
        terms = [t for t in query_lower.replace('\n', ' ').split() if len(t) > 2]
        # Add domain-specific boosts
        boosts = {
            '–¥–µ—Ç–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω': 5,
            '–æ—Ç—á–µ—Ç': 3,
            '—Ñ–æ—Ä–º–∞—Ç': 3,
            '—Ç—Ä–∞—Ñ–∏–∫': 3,
            # Domain-specific terms for billing rules
            '–∞–±–æ–Ω–µ–Ω—Ç—Å–∫–∞—è': 6,
            '–ø–ª–∞—Ç–∞': 4,
            '–≤–∫–ª—é—á–µ–Ω–Ω—ã–π': 5,
            '–≤–∫–ª—é—á–µ–Ω–Ω–æ–≥–æ': 5,
            '–ø—Ä–µ–≤—ã—à–∞—é—â': 4,
            '–ø—Ä–æ–ø–æ—Ä—Ü–∏–æ–Ω–∞–ª—å–Ω–æ': 8,
            '–¥–æ–ª–µ': 3,
            '–∞–∫—Ç–∏–≤–Ω—ã—Ö': 6,
            '–¥–Ω–µ–π': 3,
            'sbd': 5,
        }
        
        scored = []
        for doc in self.kb_chunks[kb_id]:
            text = doc.page_content.lower()
            score = 0
            for t in terms:
                cnt = text.count(t)
                if cnt:
                    score += cnt * boosts.get(t, 1)
            # phrase boost
            if '–¥–µ—Ç–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω–æ–≥–æ –æ—Ç—á–µ—Ç–∞' in text:
                score += 20
            if score > 0:
                new_meta = dict(doc.metadata)
                new_meta['kb_name'] = self.kb_metadata[kb_id]['name']
                new_meta['kb_category'] = self.kb_metadata[kb_id]['category']
                new_meta['search_type'] = 'keyword_match'
                scored.append((score, Document(page_content=doc.page_content, metadata=new_meta)))
        
        scored.sort(key=lambda x: x[0], reverse=True)
        return [d for _, d in scored[:k]]
    
    def _deduplicate_documents(self, documents: List[Document]) -> List[Document]:
        """Remove duplicate documents based on content"""
        seen_content = set()
        unique_docs = []
        
        for doc in documents:
            content_hash = hash(doc.page_content[:100])  # Use first 100 chars as identifier
            if content_hash not in seen_content:
                seen_content.add(content_hash)
                unique_docs.append(doc)
        
        return unique_docs
    
    def get_response_with_context(self, question: str, kb_ids: List[int] = None, 
                                 context_limit: int = 3) -> str:
        """Get response using RAG with specific knowledge bases"""
        try:
            # Search for relevant documents
            relevant_docs = self.search_across_kbs(question, kb_ids, k=context_limit)
            
            if not relevant_docs:
                return "–ù–µ –Ω–∞–π–¥–µ–Ω–æ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ–π –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –≤ –±–∞–∑–∞—Ö –∑–Ω–∞–Ω–∏–π."
            
            # Format context
            context = self._format_documents(relevant_docs)
            
            # Create prompt
            template = """
–¢—ã ‚Äî –∞—Å—Å–∏—Å—Ç–µ–Ω—Ç –±–∏–ª–ª–∏–Ω–≥–∞. –û—Ç–≤–µ—á–∞–π —Å—Ç—Ä–æ–≥–æ –ü–û –ö–û–ù–¢–ï–ö–°–¢–£ –∏–∑ –±–∞–∑—ã –∑–Ω–∞–Ω–∏–π.

–ö–û–ù–¢–ï–ö–°–¢:
{context}

–í–û–ü–†–û–°: {question}

–ü–†–ê–í–ò–õ–ê –û–¢–í–ï–¢–ê (–û–ë–Ø–ó–ê–¢–ï–õ–¨–ù–´):
1) –¢–û–õ–¨–ö–û –Ω–∞ —Ä—É—Å—Å–∫–æ–º —è–∑—ã–∫–µ.
2) –ò—Å–ø–æ–ª—å–∑—É–π –¢–û–õ–¨–ö–û —Å–≤–µ–¥–µ–Ω–∏—è –∏–∑ –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞: —Ñ–æ—Ä–º—É–ª–∏—Ä–æ–≤–∫–∏, –ø—É–Ω–∫—Ç—ã, –ø—Ä–∞–≤–∏–ª–∞.
3) –ï—Å–ª–∏ –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ –µ—Å—Ç—å —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è, –∏—Å–ø–æ–ª—å–∑—É–π –µ—ë –¥–ª—è –æ—Ç–≤–µ—Ç–∞, –¥–∞–∂–µ –µ—Å–ª–∏ –æ–Ω–∞ –Ω–µ –ø–æ–ª–Ω–∞—è.
4) –ï—Å–ª–∏ –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ –ù–ï–¢ –¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ–π –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏ –¥–ª—è –ø–æ–ª–Ω–æ–≥–æ –æ—Ç–≤–µ—Ç–∞, –Ω–æ –µ—Å—Ç—å —á–∞—Å—Ç–∏—á–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è, –æ—Ç–≤–µ—Ç—å –Ω–∞ –æ—Å–Ω–æ–≤–µ –∏–º–µ—é—â–∏—Ö—Å—è –¥–∞–Ω–Ω—ã—Ö –∏ —É–∫–∞–∂–∏, —á—Ç–æ –¥–ª—è –ø–æ–ª–Ω–æ–≥–æ –æ—Ç–≤–µ—Ç–∞ –Ω—É–∂–Ω–∞ –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è.
5) –¢–æ–ª—å–∫–æ –µ—Å–ª–∏ –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç–µ –≤–æ–æ–±—â–µ –ù–ï–¢ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ–π –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏, –æ—Ç–≤–µ—Ç—å: "–ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ –¥–∞–Ω–Ω—ã—Ö –≤ –ë–ó –¥–ª—è —Ç–æ—á–Ω–æ–≥–æ –æ—Ç–≤–µ—Ç–∞".
6) –î–ª—è —Ä–∞—Å—á—ë—Ç–Ω—ã—Ö –≤–æ–ø—Ä–æ—Å–æ–≤ –¥–∞–π –∫—Ä–∞—Ç–∫—É—é —Ñ–æ—Ä–º—É–ª—É/—à–∞–≥–∏ —Å—Ç—Ä–æ–≥–æ –ø–æ —Ç–µ–∫—Å—Ç—É –ë–ó.
7) –ù–µ –ø—Ä–∏–¥—É–º—ã–≤–∞–π —Ñ–∞–∫—Ç—ã –≤–Ω–µ –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞.

–û–¢–í–ï–¢ (–∫—Ä–∞—Ç–∫–æ, –ø–æ –¥–µ–ª—É):
"""

            prompt = ChatPromptTemplate.from_template(template)
            # Keep model output as AIMessage to access usage metadata when available
            chain = prompt | self.chat_model
            
            # Simple direct call without timeout for now
            # TODO: Implement proper timeout mechanism
            model_output = chain.invoke({
                "context": context,
                "question": question
            })
            
            # Extract text content
            response_text = getattr(model_output, 'content', None) or str(model_output)
            
            # Try to capture usage metadata (supported by ChatOpenAI backends)
            usage = getattr(model_output, 'usage_metadata', None)
            if usage is None:
                # Some integrations place usage under response_metadata
                metadata = getattr(model_output, 'response_metadata', {}) or {}
                try:
                    usage = (metadata or {}).get('token_usage') or (metadata or {}).get('usage')
                except Exception:
                    usage = None
            
            # Log usage if any
            try:
                self._log_llm_usage(
                    provider=self._chat_backend.get('provider'),
                    model=self._chat_backend.get('model'),
                    prompt_tokens=(usage or {}).get('prompt_tokens'),
                    completion_tokens=(usage or {}).get('completion_tokens'),
                    total_tokens=(usage or {}).get('total_tokens'),
                    question=question,
                    response_length=len(response_text or ""),
                )
            except Exception:
                pass
            
            return response_text
            
        except Exception as e:
            return f"–û—à–∏–±–∫–∞ –ø–æ–ª—É—á–µ–Ω–∏—è –æ—Ç–≤–µ—Ç–∞: {e}"
    
    def preview_context(self, question: str, kb_ids: List[int] = None, k: int = 4) -> Dict:
        """–í–µ—Ä–Ω—É—Ç—å —Å—Ñ–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–Ω—ã–π –∫–æ–Ω—Ç–µ–∫—Å—Ç –∏ –∏—Å—Ç–æ—á–Ω–∏–∫–∏ –¥–ª—è –æ—Ç–ª–∞–¥–∫–∏/–æ—Ç–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ UI."""
        try:
            relevant_docs = self.search_across_kbs(question, kb_ids, k=k)
            if not relevant_docs:
                return {"context": "", "sources": [], "docs": []}
            context = self._format_documents(relevant_docs)
            sources = []
            for doc in relevant_docs:
                _meta = doc.metadata if isinstance(doc.metadata, dict) else {}
                sources.append({
                    "preview": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                    "kb_name": _meta.get('kb_name', 'Unknown'),
                    "kb_category": _meta.get('kb_category', 'Unknown'),
                    "title": _meta.get('title', '')
                })
            return {"context": context, "sources": sources, "docs": relevant_docs}
        except Exception:
            return {"context": "", "sources": [], "docs": []}
    
    def _format_documents(self, docs: List[Document]) -> str:
        """Format documents for context"""
        formatted = []
        for i, doc in enumerate(docs, 1):
            _meta = doc.metadata if isinstance(doc.metadata, dict) else {}
            kb_name = _meta.get('kb_name', '–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –ë–ó')
            title = _meta.get('title', '–ë–µ–∑ –Ω–∞–∑–≤–∞–Ω–∏—è')
            search_type = _meta.get('search_type', 'vector_search')
            
            # Show more content for better context
            content = doc.page_content[:800] if len(doc.page_content) > 800 else doc.page_content
            
            formatted.append(f"""
–î–æ–∫—É–º–µ–Ω—Ç {i} (–∏–∑ –ë–ó: {kb_name}):
–ù–∞–∑–≤–∞–Ω–∏–µ: {title}
–¢–∏–ø –ø–æ–∏—Å–∫–∞: {search_type}
–°–æ–¥–µ—Ä–∂–∞–Ω–∏–µ: {content}
""")
        
        return "\n".join(formatted)
    
    def get_kb_statistics(self) -> Dict:
        """Get statistics about loaded knowledge bases"""
        total_docs = sum(metadata['doc_count'] for metadata in self.kb_metadata.values())
        total_chunks = sum(metadata['chunk_count'] for metadata in self.kb_metadata.values())
        
        return {
            'loaded_kbs': len(self.vectorstores),
            'total_documents': total_docs,
            'total_chunks': total_chunks,
            'kb_details': self.kb_metadata
        }
    
    def reload_kb(self, kb_id: int) -> bool:
        """Reload specific knowledge base"""
        if kb_id in self.vectorstores:
            del self.vectorstores[kb_id]
        if kb_id in self.kb_metadata:
            del self.kb_metadata[kb_id]
        
        return self.load_knowledge_base(kb_id)
    
    def clear_all(self):
        """Clear all loaded knowledge bases"""
        self.vectorstores.clear()
        self.kb_metadata.clear()
    
    # API wrapper methods
    def get_status(self) -> Dict:
        """Get RAG system status for API"""
        try:
            loaded_kbs = list(self.vectorstores.keys())
            available_kbs = self.get_available_kbs()
            stats = self.get_kb_statistics()
            
            return {
                "status": "ready" if loaded_kbs else "no_kbs_loaded",
                "loaded_kbs": loaded_kbs,
                "available_kbs": available_kbs,
                "statistics": stats,
                "embedding_backend": self._embedding_backend,
                "chat_backend": self._chat_backend
            }
        except Exception as e:
            return {
                "status": "error",
                "message": str(e),
                "loaded_kbs": [],
                "available_kbs": []
            }
    
    def search(self, query: str, kb_names: Optional[List[str]] = None, limit: int = 5) -> List[Dict]:
        """Search in knowledge bases - API wrapper"""
        try:
            # Convert kb_names to kb_ids if provided
            kb_ids = None
            if kb_names:
                kb_ids = []
                for kb_name in kb_names:
                    for kb_id, metadata in self.kb_metadata.items():
                        if metadata.get('name') == kb_name:
                            kb_ids.append(kb_id)
                            break
            
            # Search across knowledge bases
            docs = self.search_across_kbs(query, kb_ids, k=limit)
            
            # Convert to API format
            results = []
            for doc in docs:
                results.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "kb_id": doc.metadata.get('kb_id', 'Unknown'),
                    "kb_name": doc.metadata.get('kb_name', 'Unknown'),
                    "kb_category": doc.metadata.get('kb_category', 'Unknown')
                })
            
            return results
        except Exception as e:
            return [{"error": str(e)}]
    
    def ask_question(self, question: str, kb_names: Optional[List[str]] = None, context_limit: int = 3) -> Dict:
        """Ask a question using RAG - API wrapper"""
        try:
            # Convert kb_names to kb_ids if provided
            kb_ids = None
            if kb_names:
                kb_ids = []
                for kb_name in kb_names:
                    for kb_id, metadata in self.kb_metadata.items():
                        if metadata.get('name') == kb_name:
                            kb_ids.append(kb_id)
                            break
            
            # Get response with context
            answer = self.get_response_with_context(question, kb_ids, context_limit=context_limit)
            
            # Get sources
            relevant_docs = self.search_across_kbs(question, kb_ids, k=context_limit)
            sources = []
            kb_used = set()
            
            for doc in relevant_docs:
                _meta = doc.metadata if isinstance(doc.metadata, dict) else {}
                sources.append({
                    "content": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                    "kb_name": _meta.get('kb_name', 'Unknown'),
                    "kb_category": _meta.get('kb_category', 'Unknown')
                })
                kb_used.add(_meta.get('kb_name', 'Unknown'))
            
            return {
                "answer": answer,
                "sources": sources,
                "kb_used": list(kb_used)
            }
        except Exception as e:
            return {
                "answer": f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ –≤–æ–ø—Ä–æ—Å–∞: {str(e)}",
                "sources": [],
                "kb_used": []
            }
    
    def get_images_from_results(self, docs: List[Document]) -> List[Dict]:
        """–ò–∑–≤–ª–µ—á—å –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏–∑ —Ä–µ–∑—É–ª—å—Ç–∞—Ç–æ–≤ –ø–æ–∏—Å–∫–∞"""
        images = []
        for doc in docs:
            _meta = doc.metadata if isinstance(doc.metadata, dict) else {}
            if _meta.get('content_type') == 'image':
                images.append({
                    'image_path': _meta.get('source'),
                    'image_name': _meta.get('image_name'),
                    'image_description': _meta.get('image_description'),
                    'llava_analysis': _meta.get('llava_analysis'),
                    'kb_name': _meta.get('kb_name'),
                    'kb_category': _meta.get('kb_category'),
                    'content': doc.page_content
                })
        return images
    
    def search_with_images(self, query: str, kb_names: Optional[List[str]] = None, 
                          k: int = 5) -> Dict:
        """–ü–æ–∏—Å–∫ —Å –≤–æ–∑–≤—Ä–∞—Ç–æ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–π"""
        try:
            # –í—ã–ø–æ–ª–Ω—è–µ–º –æ–±—ã—á–Ω—ã–π –ø–æ–∏—Å–∫
            docs = self.search_across_kbs(query, kb_names, k=k)
            
            # –†–∞–∑–¥–µ–ª—è–µ–º –Ω–∞ —Ç–µ–∫—Å—Ç–æ–≤—ã–µ –¥–æ–∫—É–º–µ–Ω—Ç—ã –∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
            text_docs = [doc for doc in docs if (doc.metadata if isinstance(doc.metadata, dict) else {}).get('content_type') != 'image']
            image_docs = [doc for doc in docs if (doc.metadata if isinstance(doc.metadata, dict) else {}).get('content_type') == 'image']
            
            # –§–æ—Ä–º–∏—Ä—É–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã
            results = {
                'text_results': [],
                'image_results': [],
                'total_found': len(docs),
                'text_count': len(text_docs),
                'image_count': len(image_docs)
            }
            
            # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º —Ç–µ–∫—Å—Ç–æ–≤—ã–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã
            for doc in text_docs:
                _meta = doc.metadata if isinstance(doc.metadata, dict) else {}
                results['text_results'].append({
                    "content": doc.page_content,
                    "metadata": _meta,
                    "kb_name": _meta.get('kb_name', 'Unknown'),
                    "kb_category": _meta.get('kb_category', 'Unknown')
                })
            
            # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
            for doc in image_docs:
                _meta = doc.metadata if isinstance(doc.metadata, dict) else {}
                results['image_results'].append({
                    "image_path": _meta.get('source'),
                    "image_name": _meta.get('image_name'),
                    "image_description": _meta.get('image_description'),
                    "llava_analysis": _meta.get('llava_analysis'),
                    "kb_name": _meta.get('kb_name', 'Unknown'),
                    "kb_category": _meta.get('kb_category', 'Unknown'),
                    "content": doc.page_content
                })
            
            return results
            
        except Exception as e:
            return {
                'text_results': [],
                'image_results': [],
                'total_found': 0,
                'text_count': 0,
                'image_count': 0,
                'error': str(e)
            }